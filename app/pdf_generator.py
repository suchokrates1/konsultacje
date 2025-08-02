"""Utilities for rendering session details into a PDF document using a DOCX template."""

import os
import tempfile
import subprocess

from flask import current_app
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert


def generate_pdf(zajecia, beneficjenci, output_path):
    """Create a PDF summary for a session and write it to *output_path*.

    The function fills the ``wzor.docx`` template using :mod:`python-docx`
    and then converts the rendered document to PDF using ``docx2pdf``.
    """

    template_path = os.path.join(current_app.root_path, "static", "wzor.docx")
    if not os.path.exists(template_path):
        current_app.logger.error("Missing DOCX template: %s", template_path)
        raise FileNotFoundError(f"Template file not found: {template_path}")

    doc = Document(template_path)
    start_time = zajecia.godzina_od.strftime("%H:%M")
    end_time = zajecia.godzina_do.strftime("%H:%M")
    names = "\n".join(b.imie for b in beneficjenci[:3])
    wojew = ", ".join({b.wojewodztwo for b in beneficjenci[:3]})
    context = {
        "data": zajecia.data.strftime("%d.%m.%Y"),
        "time_range": f"{start_time} - {end_time}",
        "beneficjenci": names,
        "wojewodztwo": wojew,
        "specjalista": zajecia.specjalista,
    }
    current_app.config["_last_pdf_context"] = context

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Imię i nazwisko beneficjenta:"):
            paragraph.text = "Imię i nazwisko beneficjenta: " + names
        if text.startswith("Województwo:"):
            paragraph.text = "Województwo: " + wojew

    if doc.tables:
        table = doc.tables[0]
        for idx, _ in enumerate(beneficjenci[:3]):
            row = table.rows[idx + 1]
            values = [context["data"], context["time_range"], context["specjalista"]]
            for col, value in enumerate(values, start=1):
                cell = row.cells[col]
                cell.text = value
                if cell.paragraphs:
                    para = cell.paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if para.runs:
                        run = para.runs[0]
                        run.font.size = Pt(18)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
    try:
        doc.save(tmp_path)
        try:
            convert(tmp_path, output_path)
        except NotImplementedError:
            out_dir = os.path.dirname(output_path) or "."
            cmd = [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                out_dir,
                tmp_path,
            ]
            try:
                subprocess.run(cmd, check=True)
            except FileNotFoundError:
                cmd[0] = "soffice"
                subprocess.run(cmd, check=True)
            generated = os.path.join(
                out_dir,
                os.path.splitext(os.path.basename(tmp_path))[0] + ".pdf",
            )
            os.replace(generated, output_path)
    finally:
        try:
            os.remove(tmp_path)
        except OSError:
            current_app.logger.warning(
                "Failed to remove temporary file %s", tmp_path
            )

