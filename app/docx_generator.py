"""Utilities for rendering session details into a DOCX document."""

import os
import io

from flask import current_app
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt


def generate_docx(zajecia, beneficjenci, output_path):
    """Create a DOCX summary for a session and write it to *output_path*."""

    template_path = os.path.join(current_app.root_path, "static", "wzor.docx")
    if not os.path.exists(template_path):
        current_app.logger.error("Missing DOCX template: %s", template_path)
        raise FileNotFoundError(f"Template file not found: {template_path}")

    # Load the template into memory to avoid mutating the file on disk.
    with open(template_path, "rb") as f:
        doc = Document(io.BytesIO(f.read()))
    start_time = zajecia.godzina_od.strftime("%H:%M")
    end_time = zajecia.godzina_do.strftime("%H:%M")
    names = "\n".join(b.imie for b in beneficjenci[:3])
    # Using dict.fromkeys to keep the order of provinces aligned with
    # the order of beneficiaries while removing duplicates.
    wojewodztwa = dict.fromkeys(b.wojewodztwo for b in beneficjenci[:3])
    wojew = ", ".join(wojewodztwa)
    context = {
        "data": zajecia.data.strftime("%d.%m.%Y"),
        "time_range": f"{start_time} - {end_time}",
        "beneficjenci": names,
        "wojewodztwo": wojew,
        # full name of the instructor for later assertions/tests
        "specjalista": zajecia.user.full_name,
    }
    current_app.config["_last_docx_context"] = context

    # Replace occurrences of "dietetykiem" throughout the document.
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.text = run.text.replace("dietetykiem", zajecia.specjalista)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = run.text.replace("dietetykiem", zajecia.specjalista)

    specialist_found = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Imię i nazwisko beneficjenta:"):
            paragraph.text = "Imię i nazwisko beneficjenta: " + names
        if text.startswith("Województwo:"):
            paragraph.text = "Województwo: " + wojew
        if text.startswith("Imię i nazwisko specjalisty:"):
            paragraph.text = (
                "Imię i nazwisko specjalisty: " + zajecia.user.full_name
            )
            specialist_found = True

    if not specialist_found:
        doc.add_paragraph(
            f"Imię i nazwisko specjalisty: {zajecia.user.full_name}"
        )

    if doc.tables:
        table = doc.tables[0]
        font_size = Pt(16)
        for idx, _ in enumerate(beneficjenci[:3]):
            row = table.rows[idx + 1]
            values = [context["data"], context["time_range"], zajecia.user.full_name]
            for cell, value in zip(row.cells[1:4], values):
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                p.clear()
                run = p.add_run(value)
                run.font.size = font_size
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)
