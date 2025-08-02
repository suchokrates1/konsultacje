"""Tests for beneficiary CRUD operations and DOCX generation."""

from datetime import date, time
import io
import os

from app import db
from app.models import User, Beneficjent, Zajecia
from docx import Document


def _docx_text(doc: Document) -> str:
    """Return all text content from a DOCX document."""

    text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += "\n" + cell.text
    return text


def _docx_text_from_bytes(data: bytes) -> str:
    """Load DOCX bytes and return their textual content."""

    return _docx_text(Document(io.BytesIO(data)))


def create_user(app):
    """Create a user in the database and return its ID."""

    with app.app_context():
        user = User(full_name='tester', email='tester@example.com')
        user.set_password('secret')
        user.confirmed = True
        db.session.add(user)
        db.session.commit()
        return user.id


def login(client, username='tester', password='secret'):
    """Log in a user using the test client."""

    return client.post(
        '/login',
        data={'full_name': username, 'password': password},
        follow_redirects=True,
    )


def test_beneficjent_crud(app, client):
    """Add, edit and delete a beneficiary through the web interface."""

    create_user(app)
    login(client)

    # Verify dropdown renders
    resp = client.get('/beneficjenci/nowy')
    html = resp.get_data(as_text=True)
    assert '<select' in html
    assert html.count('<option') >= 16

    # Add
    response = client.post(
        '/beneficjenci/nowy',
        data={'imie': 'Jan Kowalski', 'wojewodztwo': 'Mazowieckie'},
        follow_redirects=True,
    )
    assert 'Beneficjent dodany' in response.get_data(as_text=True)
    with app.app_context():
        benef = Beneficjent.query.filter_by(imie='Jan Kowalski').first()
        assert benef is not None
        benef_id = benef.id

    # Edit
    response = client.post(
        f'/beneficjenci/{benef_id}/edytuj',
        data={'imie': 'Jan Nowak', 'wojewodztwo': 'Slaskie'},
        follow_redirects=True,
    )
    assert 'Beneficjent zaktualizowany' in response.get_data(as_text=True)
    with app.app_context():
        benef = db.session.get(Beneficjent, benef_id)
        assert benef.imie == 'Jan Nowak'
        assert benef.wojewodztwo == 'Slaskie'

    # Delete
    response = client.post(
        f'/beneficjenci/{benef_id}/usun',
        data={'submit': '1'},
        follow_redirects=True,
    )
    assert 'Beneficjent usunięty' in response.get_data(as_text=True)
    with app.app_context():
        assert db.session.get(Beneficjent, benef_id) is None


def test_create_session(app, client):
    """Create a consultation session and verify it was stored."""
    user_id = create_user(app)
    login(client)
    with app.app_context():
        benef = Beneficjent(
            imie='Anna', wojewodztwo='Pomorskie', user_id=user_id
        )
        db.session.add(benef)
        db.session.commit()
        b_id = benef.id

    response = client.post(
        '/zajecia/nowe',
        data={
            'data': '2023-01-01',
            'godzina_od': '10:00',
            'godzina_do': '11:00',
            'beneficjenci': str(b_id),
        },
        follow_redirects=True,
    )
    assert response.request.path == '/zajecia'
    assert 'Zajęcia zapisane' in response.get_data(as_text=True)
    with app.app_context():
        zajecia = Zajecia.query.filter_by(user_id=user_id).first()
        assert zajecia is not None
        assert len(zajecia.beneficjenci) == 1


def test_docx_generation(app, client):
    """Generate a DOCX report and ensure the file is removed afterwards."""
    user_id = create_user(app)
    login(client)
    with app.app_context():
        benef = Beneficjent(
            imie='Piotr', wojewodztwo='Lubelskie', user_id=user_id
        )
        db.session.add(benef)
        zajecia = Zajecia(
            data=date(2023, 1, 2),
            godzina_od=time(9, 0),
            godzina_do=time(10, 0),
            specjalista='tester',
            user_id=user_id
        )
        zajecia.beneficjenci.append(benef)
        db.session.add(zajecia)
        db.session.commit()
        z_id = zajecia.id

    response = client.get(f'/zajecia/{z_id}/docx')
    assert response.status_code == 200
    assert 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' in response.headers.get('Content-Type', '')
    disposition = response.headers.get('Content-Disposition', '')
    expected_name = (
        "Konsultacje dietetyczne 2023-01-02 Piotr.docx"
    )
    assert disposition.startswith('attachment')
    assert expected_name in disposition

    docx_path = os.path.join(
        app.root_path,
        'static',
        'docx',
        expected_name,
    )
    # the generated DOCX should be removed after the response is sent
    assert not os.path.exists(docx_path)


def test_docx_content(app, client):
    """Generate a DOCX and verify expected text appears in the document."""
    user_id = create_user(app)
    login(client)
    with app.app_context():
        benef = Beneficjent(
            imie="Katarzyna", wojewodztwo="Mazowieckie", user_id=user_id
        )
        db.session.add(benef)
        zajecia = Zajecia(
            data=date(2023, 2, 1),
            godzina_od=time(8, 0),
            godzina_do=time(9, 0),
            specjalista="tester",
            user_id=user_id,
        )
        zajecia.beneficjenci.append(benef)
        db.session.add(zajecia)
        db.session.commit()
        z_id = zajecia.id

    response = client.get(f"/zajecia/{z_id}/docx")
    assert response.status_code == 200

    text = _docx_text_from_bytes(response.data)
    assert "tester" in text
    assert "Katarzyna" in text
    assert "Mazowieckie" in text

