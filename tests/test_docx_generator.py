"""Tests ensuring generated DOCX files contain the expected text."""

import os
from datetime import date, time
from docx import Document

from app import db
from app.models import User, Beneficjent, Zajecia
from app.docx_generator import generate_docx


def _docx_text(doc: Document) -> str:
    """Extract all text content from a DOCX document."""

    text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += "\n" + cell.text
    return text


def test_generate_docx_file_contains_text(app, tmp_path):
    """DOCX generated on disk should include rendered session data."""

    with app.app_context():
        user = User(full_name="Jan Kowalski", email="disk@example.com")
        user.set_password("pass")
        user.confirmed = True
        db.session.add(user)
        db.session.commit()

        benef = Beneficjent(imie="Tomek", wojewodztwo="Mazowieckie", user_id=user.id)
        db.session.add(benef)
        zajecia = Zajecia(
            data=date(2023, 3, 3),
            godzina_od=time(10, 0),
            godzina_do=time(11, 0),
            specjalista="psycholog",
            user_id=user.id,
        )
        zajecia.beneficjenci.append(benef)
        db.session.add(zajecia)
        db.session.commit()

        path = tmp_path / "report.docx"
        generate_docx(zajecia, [benef], str(path))
        doc = Document(str(path))
        text = _docx_text(doc)
        assert zajecia.specjalista in text
        assert user.full_name in text
        assert "Tomek" in text
        assert "dietetykiem" not in text
        # The table column labeled "Imię i nazwisko specjalisty" should contain
        # the instructor's full name for each listed session.
        table = doc.tables[0]
        assert table.rows[1].cells[3].text == user.full_name

        template = Document(os.path.join(app.root_path, "static", "wzor.docx"))
        assert "dietetykiem" in _docx_text(template)


def test_generate_docx_adds_missing_rows(app, tmp_path, monkeypatch):
    """Generator should expand table when template lacks enough rows."""

    with app.app_context():
        template_dir = tmp_path / "tmpl"
        static_dir = template_dir / "static"
        static_dir.mkdir(parents=True)
        template_path = static_dir / "wzor.docx"

        doc = Document()
        doc.add_table(rows=1, cols=4)
        doc.save(str(template_path))

        monkeypatch.setattr(app, "root_path", str(template_dir))

        user = User(full_name="Jan Kowalski", email="dyn@example.com")
        user.set_password("pass")
        user.confirmed = True
        db.session.add(user)
        db.session.commit()

        b1 = Beneficjent(imie="Anna", wojewodztwo="Mazowieckie", user_id=user.id)
        b2 = Beneficjent(imie="Ewa", wojewodztwo="Mazowieckie", user_id=user.id)
        db.session.add_all([b1, b2])

        zajecia = Zajecia(
            data=date(2023, 3, 3),
            godzina_od=time(10, 0),
            godzina_do=time(11, 0),
            specjalista="psycholog",
            user_id=user.id,
        )
        zajecia.beneficjenci.extend([b1, b2])
        db.session.add(zajecia)
        db.session.commit()

        output = tmp_path / "report.docx"
        generate_docx(zajecia, [b1, b2], str(output))

        out_doc = Document(str(output))
        table = out_doc.tables[0]
        assert len(table.rows) == 3
        assert table.rows[2].cells[3].text == user.full_name
